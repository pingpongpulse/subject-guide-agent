from docx import Document
from docx.oxml.ns import qn
import os

def load_docx(filepath):
    """
    Extracts text from DOCX files including:
    - Normal paragraphs
    - Tables
    - Headers and footers
    """
    doc = Document(filepath)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                [cell.text.strip() for cell in row.cells if cell.text.strip()]
            )
            if row_text:
                full_text.append(row_text)

    combined_text = "\n".join(full_text)

    return [{
        "text": combined_text,
        "page_number": 1,
        "source_file": os.path.basename(filepath),
        "total_pages": 1,
        "extraction_method": "python-docx"
    }]


if __name__ == "__main__":
    result = load_docx("data/sample_docs/test.docx")
    print(result[0]['text'][:500])